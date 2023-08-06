#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from __future__ import division
import sys
import os
import re
import json
import codecs
import requests
import pdb
import webbrowser
from collections import defaultdict
from chgksuite.common import (
    get_lastdir, set_lastdir, on_close, log_wrap, bring_to_front,
    get_chgksuite_dir
)
try:
    basestring
except NameError:
    basestring = str
try:
    input = raw_input
except NameError:
    pass


API = 'https://trello.com/1'
re_bi = re.compile(r'trello\.com/b/(.+?)(/|$)')



def upload_file(filepath, trello):
    req = requests.get("{}/boards/{}/lists".format(API, trello['board_id']),
                       params={'token': trello['params']['token'],
                               'key': trello['params']['key']})
    if req.status_code != 200:
        print('Error: {}'.format(req.text))
        sys.exit(1)

    lists = json.loads(req.content.decode('utf8'))
    lid = lists[0]['id']
    content = ''
    with codecs.open(filepath, 'r', 'utf8') as f:
        content = f.read()
    cards = re.split(r'(\r?\n){2,}', content)
    cards = [x for x in cards if x != '' and x != '\n' and x != '\r\n']
    for card in cards:
        caption = 'вопрос'
        if re.search('\n! (.+?)\r?\n', card):
            caption = re.search('\n! (.+?)\.?\r?\n', card).group(1)
            if trello['author'] and re.search('\n@ (.+?)\.?\r?\n', card):
                caption += ' {}'.format(
                    re.search('\n@ (.+?)\r?\n', card).group(1))

        req = requests.post(
            "{}/lists/{}/cards".format(API, lid),
            {
                'key': trello['params']['key'],
                'token': trello['params']['token'],
                'desc': card,
                'name': caption
            })
        if req.status_code == 200:
            print('Successfully sent {}'.format(log_wrap(caption)))
        else:
            print('Error {}: {}'.format(req.status_code, req.content))


def gui_trello_upload(args, sourcedir):
    ld = get_lastdir()

    if not args.board_id:
        board_id = get_board_id()
    else:
        board_id = args.board_id

    trelloconfig = args.trelloconfig
    trelloconfig['board_id'] = board_id

    trelloconfig['author'] = args.author

    if isinstance(args.filename, (list, tuple)):
        if len(args.filename) == 1 and os.path.isdir(args.filename[0]):
            for filename in os.listdir(args.filename[0]):
                if filename.endswith('.4s'):
                    filepath = os.path.join(args.filename[0], filename)
                    upload_file(filepath, trelloconfig)
            set_lastdir(args.filename[0])
        else:
            for filename in args.filename:
                upload_file(filename, trelloconfig)
                set_lastdir(filename)
    elif isinstance(args.filename, basestring):
        if os.path.isdir(args.filename):
            for filename in os.listdir(args.filename):
                if filename.endswith('.4s'):
                    filepath = os.path.join(args.filename, filename)
                    upload_file(filepath, trelloconfig)
                    set_lastdir(filepath)
        elif os.path.isfile(args.filename):
            upload_file(args.filename, trelloconfig)
            set_lastdir(args.filename)


def onlyanswers_line_check(line):
    line = (line or '')
    return line.startswith(
        ('Ответ', 'Зачёт', 'Зачет', '1', '2', '3', '4', '5', '6', '8')
    )


def noanswers_line_check(line):
    line = (line or '')
    return not line.startswith(
        (
            'Ответ', 'Коммента', 'Источник', 'Автор', 'Зачёт', 'Зачет',
            'Незачёт', 'Незачет'
        )
    )


def process_desc(s, onlyanswers=False, noanswers=False):
    s = s.strip()
    s = s.replace(r'\`', '`')
    s = s.replace(r'\*', '*')
    if onlyanswers:
        lines = s.split('\n')
        lines = [
            x for x in lines if onlyanswers_line_check(x)
        ]
        s = '\n'.join(lines)
    elif noanswers:
        lines = s.split('\n')
        lines = [
            x for x in lines if noanswers_line_check(x)
        ]
        s = '\n'.join(lines)
    return s


def getlabels(s):
    return {x['name'] for x in s['labels']}


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def add_themes_list(group):
    themes = group['themes']
    p = group['paragraph']
    if len(themes):
        p.add_run('Темы: ' + ', '.join(themes) + '.').bold = True
    else:
        delete_paragraph(p)


def gui_trello_download(args, sourcedir):
    ld = get_lastdir()

    template_path = args.docx_template

    board_id_path = os.path.join(args.folder, '.board_id')
    if os.path.isfile(board_id_path):
        with codecs.open(board_id_path, 'r', 'utf8') as f:
            board_id = f.read().rstrip()
    else:
        board_id = get_board_id(path=args.folder)

    params = args.trelloconfig['params']
    ld = args.folder
    set_lastdir(ld)
    os.chdir(args.folder)

    if args.si or args.qb:
        from docx import Document

    req = requests.get("{}/boards/{}".format(API, board_id),
                       params=params)
    if req.status_code != 200:
        print('Error: {}'.format(req.text))
        if args.debug:
            pdb.set_trace()
        sys.exit(1)

    _lists = defaultdict(lambda: [])
    _list_counters = defaultdict(lambda: 0)

    json_ = json.loads(req.content.decode('utf8'))
    _names = defaultdict(lambda: None)
    open_lists = list(filter(lambda l: not l['closed'], json_['lists']))
    for list_ in open_lists:
        _names[list_['id']] = list_['name'].replace('/', '_')
        _list_counters[list_['id']] = 0

    if args.si:
        _docs = defaultdict(lambda: Document(template_path))
        _groups = defaultdict(lambda: None)
    if args.qb:
        qb_doc = Document(template_path)

    for card in json_['cards']:
        card["desc"] = card["desc"].replace("\n\n", "\n")
        list_id = card['idList']
        list_name = _names[list_id]
        if card.get('closed') or list_name is None:
            continue

        _list_counters[list_id] += 1

        if not args.si:
            card_title = ''
        elif card['name'].startswith('#'):
            card_title = card['name']
            _list_counters[list_id] = 0
        else:
            card_title = 'Тема {}. {}'.format(
                _list_counters[list_id], card['name']
            )
            clear_card_title = card['name']

        id_ = ('singlefile' if args.singlefile else list_name)

        if args.si:
            doc_ = _docs[id_]
            group_ = _groups[id_]
            if group_ is None:  # new doc
                group_ = _groups[id_] = {
                    'paragraph': doc_.add_paragraph(),
                    'themes': []
                }
            if card_title:  # new title
                if card_title.startswith('#'):
                    title_re = r'(#+)\s*(.*)'
                    m = re.search(title_re, card_title)
                    doc_.add_heading(m[2], level=len(m[1]))
                    add_themes_list(group_)
                    group_['paragraph'] = doc_.add_paragraph()
                    group_['themes'] = []
                    doc_.add_paragraph()
                else:
                    group_['themes'].append(clear_card_title)
                    p = doc_.add_paragraph()
                    p.add_run(card_title).bold = True
                    doc_.add_paragraph()
            if card['desc']:
                doc_.add_paragraph(process_desc(
                    card['desc'],
                    onlyanswers=args.onlyanswers,
                    noanswers=args.noanswers,
                ))

        _lists[id_].append(card_title + ('' if card_title.startswith('#')
                                         else '\n\n') + process_desc(card['desc']))

        if args.labels:
            for label in getlabels(card):
                _lists[label].append(
                    (card['name'] if args.si else '') +
                    process_desc(card['desc']))
    if args.si:
        for list_name in _groups:
            # add remaining themes when we know we're done with the doc
            add_themes_list(_groups[list_name])
        for list_name in _docs:
            _docs[list_name].save('{}.docx'.format(list_name))

    if args.qb:
        first, second = _lists[args.qb[0]], _lists[args.qb[1]]
        for i, pair in enumerate(zip(first, second)):
            p = qb_doc.add_paragraph()
            p.add_run('Тоссап {}.'.format(i + 1)).bold = True
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run(pair[0])
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run('Бонус {}.'.format(i + 1)).bold = True
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run(pair[1])
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
        qb_doc.save('quizbowl.docx')

    for _list in _lists:
        filename = '{}.4s'.format(_list)
        print('outputting {}'.format(filename))
        with codecs.open(filename, 'w', 'utf8') as f:
            for item in _lists[_list]:
                f.write('\n' + item + '\n')


def get_board_id(path=None):
    print('To communicate with your trello board we need its board_id.')
    print('Your board link looks like this:')
    print()
    print('https://trello.com/b/Bi0z2H49/title-of-your-board')
    print('                     board_id')
    print()
    board_id = input('Please paste your board_id '
                     '(or the board link, '
                     'we\'ll parse it): ').rstrip()
    if 'trello.com' in board_id:
        board_id = re_bi.search(board_id).group(1)
    if path:
        with codecs.open(os.path.join(path, '.board_id'), 'w', 'utf8') as f:
            f.write(board_id)
    return board_id


def get_token(tokenpath):
    webbrowser.open('https://trello.com/1/connect'
                    '?key=1d4fe71dd193855686196e7768aa4b05'
                    '&name=Chgk&scope=read,write&response_type=token')
    token = input('Please paste the obtained token: ').rstrip()
    with codecs.open(tokenpath, 'w', 'utf8') as f:
        f.write(token)
    return token


def gui_trello(args, sourcedir=None):
    csdir = get_chgksuite_dir()
    resourcedir = os.path.join(sourcedir, "resources")
    tokenpath = os.path.join(csdir, '.trello_token')
    if not os.path.isfile(tokenpath):
        token = get_token(tokenpath)
    else:
        with codecs.open(tokenpath, 'r', 'utf8') as f:
            token = f.read().rstrip()

    with open(os.path.join(resourcedir, 'trello.json')) as f:
        args.trelloconfig = json.load(f)
    args.trelloconfig['params']['token'] = token

    if args.trellosubcommand == 'download':
        gui_trello_download(args, sourcedir)
    elif args.trellosubcommand == 'upload':
        gui_trello_upload(args, sourcedir)
    elif args.trellosubcommand == 'token':
        get_token(tokenpath)


def main():
    print('This program was not designed to run standalone.')
    input("Press Enter to continue...")


if __name__ == "__main__":
    main()
